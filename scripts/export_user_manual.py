from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def ensure_styles(document: Document) -> None:
    styles = document.styles

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(9.5)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)

    if "Body Tight" not in styles:
        body_style = styles.add_style("Body Tight", WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = "Aptos"
        body_style.font.size = Pt(10.5)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15


def add_inline_runs(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_CODE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        run = paragraph.add_run(match.group(1))
        run.font.name = "Courier New"
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_image(document: Document, markdown_path: Path, alt: str, target: str) -> None:
    image_path = (markdown_path.parent / target).resolve()
    if not image_path.exists():
        document.add_paragraph(f"[Missing image: {target}]", style="Body Tight")
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    width = Inches(1.2) if "logo" in alt.lower() else Inches(6.4)
    run.add_picture(str(image_path), width=width)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value.strip()
            if r_idx == 0:
                set_cell_shading(cell, "D9E2F3")


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and is_table_line(lines[i]):
        row = [part.strip() for part in lines[i].strip()[1:-1].split("|")]
        rows.append(row)
        i += 1

    if len(rows) >= 2 and all(set(part) <= {"-", ":"} for part in rows[1]):
        rows.pop(1)
    return rows, i


def export_markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    ensure_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    in_mermaid = False
    mermaid_index = 0
    mermaid_dir = markdown_path.parent / ".generated" / "user-manual-diagrams"

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if stripped == "```mermaid" and not in_code:
                in_code = True
                in_mermaid = True
                code_buffer = []
                i += 1
                continue
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                if in_mermaid:
                    mermaid_index += 1
                    image_path = mermaid_dir / f"manual-diagram-{mermaid_index:02d}.png"
                    if image_path.exists():
                        paragraph = document.add_paragraph()
                        paragraph.alignment = 1
                        run = paragraph.add_run()
                        run.add_picture(str(image_path), width=Inches(6.3))
                    else:
                        paragraph = document.add_paragraph("Mermaid diagram image missing.", style="Body Tight")
                else:
                    paragraph = document.add_paragraph(style="Code Block")
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    for idx, code_line in enumerate(code_buffer):
                        run = paragraph.add_run(code_line)
                        run.font.name = "Courier New"
                        if idx < len(code_buffer) - 1:
                            run.add_break()
                in_code = False
                in_mermaid = False
                code_buffer = []
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            add_image(document, markdown_path, image_match.group("alt"), image_match.group("path"))
            i += 1
            continue

        if is_table_line(line):
            rows, next_i = parse_table(lines, i)
            add_table(document, rows)
            i = next_i
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            paragraph = document.add_paragraph(style=f"Heading {min(level, 9)}")
            add_inline_runs(paragraph, text)
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, numbered.group(2))
            i += 1
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:].strip())
            i += 1
            continue

        paragraph = document.add_paragraph(style="Body Tight")
        add_inline_runs(paragraph, stripped)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown_path", type=Path)
    parser.add_argument("output_path", type=Path)
    args = parser.parse_args()
    export_markdown_to_docx(args.markdown_path.resolve(), args.output_path.resolve())


if __name__ == "__main__":
    main()
